import streamlit as st
from docx import Document
import datetime
from io import BytesIO
import os
import base64
import tempfile
from docx2pdf import convert
import uuid

def fill_placeholders(doc: Document, replacements: dict):
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        inline[i].text = inline[i].text.replace(key, val)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)

def generate_ili_offer(data):
    template_path = "quotation_template_ili.docx"
    doc = Document(template_path)

    replacements = {
        "<<project_name>>": data["project_name"],
        "<<client_name>>": data["client_name"],
        "<<contact_person>>": data["contact_person"],
        "<<offer_number>>": data["offer_number"],
        "<<quotation_date>>": data["quotation_date"].strftime("%d-%m-%Y"),
        "<<mobilization_days>>": str(data["mobilization_days"]),
        "<<inspection_days>>": str(data["inspection_days"]),
        "<<total_days>>": str(data["mobilization_days"] + data["inspection_days"]),
        "<<pipe_diameter>>": str(data["pipe_diameter"]),
        "<<mfl_tool_rerun>>": data["mfl_tool_rerun"],
        "<<egp_tool_rerun>>": data["egp_tool_rerun"],
        "<<egp_additional_mob>>": data["egp_additional_mob"],
        "<<mfl_additional_mob>>": data["mfl_additional_mob"],
        "<<personnel_additional_mob>>": data["personnel_additional_mob"],
        "<<mfl_standby_rate>>": data["mfl_standby_rate"],
        "<<egp_standby_rate>>": data["egp_standby_rate"],
        "<<personnel_standby_rate>>": data["personnel_standby_rate"]
    }

    for i in range(1, 8):
        replacements[f"<<segment_{i}>>"] = data["segments"].get(f"segment_{i}", "")
        replacements[f"<<price_segment_{i}>>"] = data["segments"].get(f"price_segment_{i}", "")

    fill_placeholders(doc, replacements)

    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    # Save to temp .docx file
    tmp_id = str(uuid.uuid4())
    docx_path = os.path.join(tempfile.gettempdir(), f"{tmp_id}.docx")
    pdf_path = os.path.join(tempfile.gettempdir(), f"{tmp_id}.pdf")

    with open(docx_path, "wb") as f:
        f.write(docx_buffer.read())

    # Convert DOCX to PDF
    convert(docx_path, pdf_path)

    return docx_path, pdf_path

def embed_pdf(file_path):
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode("utf-8")
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600px" type="application/pdf"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

# --- Streamlit UI ---
st.title("ILI Offer Letter Generator")

with st.form("ili_form"):
    st.header("Project & Client Info")
    project_name = st.text_input("Project Name")
    client_name = st.text_input("Client Name")
    contact_person = st.text_input("Contact Person")
    offer_number = st.text_input("Offer Number")
    quotation_date = st.date_input("Quotation Date", value=datetime.date.today())

    st.header("Inspection Details")
    mobilization_days = st.number_input("Mobilization Days", min_value=0, step=1)
    inspection_days = st.number_input("Inspection Days", min_value=0, step=1)
    pipe_diameter = st.text_input("Pipeline Diameter")

    st.header("Segment Pricing (max 7)")
    segments = {}
    for i in range(1, 8):
        segments[f"segment_{i}"] = st.text_input(f"Segment {i} Description", key=f"seg{i}")
        segments[f"price_segment_{i}"] = st.text_input(f"Segment {i} Price", key=f"price{i}")

    st.header("Reruns & Additional Charges")
    mfl_tool_rerun = st.text_input("MFL Tool Rerun Charge")
    egp_tool_rerun = st.text_input("EGP Tool Rerun Charge")
    egp_additional_mob = st.text_input("EGP Additional Mobilization Charge")
    mfl_additional_mob = st.text_input("MFL Additional Mobilization Charge")
    personnel_additional_mob = st.text_input("Personnel Additional Mobilization Charge")

    st.header("Standby Charges")
    mfl_standby_rate = st.text_input("MFL Standby Rate per Day")
    egp_standby_rate = st.text_input("EGP Standby Rate per Day")
    personnel_standby_rate = st.text_input("Personnel Standby Rate per Day")

    submitted = st.form_submit_button("Generate Offer")

if submitted:
    data = {
        "project_name": project_name,
        "client_name": client_name,
        "contact_person": contact_person,
        "offer_number": offer_number,
        "quotation_date": quotation_date,
        "mobilization_days": mobilization_days,
        "inspection_days": inspection_days,
        "pipe_diameter": pipe_diameter,
        "segments": segments,
        "mfl_tool_rerun": mfl_tool_rerun,
        "egp_tool_rerun": egp_tool_rerun,
        "egp_additional_mob": egp_additional_mob,
        "mfl_additional_mob": mfl_additional_mob,
        "personnel_additional_mob": personnel_additional_mob,
        "mfl_standby_rate": mfl_standby_rate,
        "egp_standby_rate": egp_standby_rate,
        "personnel_standby_rate": personnel_standby_rate
    }

    docx_path, pdf_path = generate_ili_offer(data)

    st.success("Offer generated successfully!")
    with open(docx_path, "rb") as f:
        st.download_button("Download Word File", f, f"{offer_number}_ILI_Offer.docx")

    embed_pdf(pdf_path)
